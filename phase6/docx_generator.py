"""
Generates a Word document (.docx) itinerary — the "PDF Generation" deliverable
from the architecture doc, produced as .docx instead of PDF per project
decision (python-docx is already installed; PDF libraries need extra
system-level setup on Windows).
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from citation_format import format_citation_label
from stop_format import format_opening_hours, format_display_name

DAY_SLOTS = (("Morning", "morning"), ("Afternoon", "afternoon"), ("Evening", "evening"))

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_line(doc: Document, line: str) -> None:
    """
    Minimal Markdown->docx line renderer for the trained narrator's output:
    '#'/'##'/'###' headings, '-'/'*' bullets, '**bold**' spans. Anything else
    becomes a plain paragraph. Not a general Markdown parser — just enough
    for the narrator's fixed output shape.
    """
    stripped = line.strip()
    if not stripped:
        return

    heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
    if heading_match:
        level = min(len(heading_match.group(1)) + 1, 4)  # narrative's "#"/"##" sit below the doc's own Day headings
        doc.add_heading(heading_match.group(2), level=level)
        return

    is_bullet = stripped.startswith(("- ", "* "))
    text = stripped[2:].strip() if is_bullet else stripped
    p = doc.add_paragraph(style="List Bullet" if is_bullet else None)

    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def build_itinerary_docx(
    itinerary: dict,
    ctx: dict,
    citations: list[dict] | None = None,
    narrative: str | None = None,
) -> io.BytesIO:
    """Returns an in-memory .docx file as a BytesIO stream."""
    doc = Document()

    title = doc.add_heading("New Delhi Travel Itinerary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_parts = [f"{ctx.get('num_days', '?')} days", f"{ctx.get('pace', 'moderate')} pace"]
    if ctx.get("interests"):
        meta_parts.append("Interests: " + ", ".join(ctx["interests"]))
    if ctx.get("group_size"):
        meta_parts.append(f"Group size: {ctx['group_size']}")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(" | ".join(meta_parts)).italic = True

    if narrative:
        for line in narrative.splitlines():
            _add_markdown_line(doc, line)
        doc.add_page_break()
        doc.add_heading("Precise Day-by-Day Schedule", level=1)
        doc.add_paragraph(
            "The grounded schedule below is the exact source data behind the overview "
            "above — real places, distances, and times from the trip planner's backend."
        ).italic = True

    day_keys = sorted((k for k in itinerary if k.startswith("day_")), key=lambda k: int(k.split("_")[1]))
    for key in day_keys:
        day = itinerary[key]
        day_num = key.split("_")[1]
        heading_text = f"Day {day_num}"
        if day.get("date"):
            heading_text += f" — {day['date']}"
        doc.add_heading(heading_text, level=1)
        doc.add_paragraph(f"Total scheduled time: {day.get('total_hours', 0)}h")

        emergency_bits = []
        if day.get("nearest_hospital"):
            h = day["nearest_hospital"]
            emergency_bits.append(f"Nearest hospital: {format_display_name(h['name'])} ({h['distance_km']} km)")
        if day.get("nearest_pharmacy"):
            p = day["nearest_pharmacy"]
            emergency_bits.append(f"Nearest pharmacy: {format_display_name(p['name'])} ({p['distance_km']} km)")
        if day.get("nearest_metro_station"):
            m = day["nearest_metro_station"]
            emergency_bits.append(f"Nearest metro: {format_display_name(m['name'])} ({m['distance_km']} km)")
        if emergency_bits:
            doc.add_paragraph(" | ".join(emergency_bits)).italic = True

        for slot_label, slot_key in DAY_SLOTS:
            stops = day.get(slot_key, [])
            if not stops:
                continue
            doc.add_heading(slot_label, level=2)
            for stop in stops:
                travel = stop.get("travel_time_from_prev_min", 0)
                mode = stop.get("travel_mode_from_prev")
                meal = stop.get("meal")
                hours = stop.get("opening_hours")
                website = stop.get("website")

                p = doc.add_paragraph(style="List Bullet")
                name_suffix = " (hidden gem)" if stop.get("is_hidden_gem") else ""
                category_suffix = f", {meal}" if meal else ""
                p.add_run(f"{format_display_name(stop['name'])}{name_suffix} ({stop['category']}{category_suffix})").bold = True
                p.add_run(f" — visit ~{stop['visit_duration_min']} min")
                if travel:
                    mode_note = f" by {mode}" if mode else ""
                    p.add_run(f", {travel} min travel{mode_note} from previous stop")
                if hours and hours != "unknown":
                    p.add_run(f" — hours: {format_opening_hours(hours)}")
                if website:
                    p.add_run(f" — {website}")
                entry_fee = stop.get("kb_entry_fee")
                if entry_fee:
                    doc.add_paragraph(f"Entry fee: {entry_fee}").italic = True

    if citations:
        doc.add_heading("Sources", level=1)
        for c in citations:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{format_citation_label(c)} — {c.get('source_url', '')}")

    doc.add_paragraph(
        "Map & POI data © OpenStreetMap contributors, ODbL (openstreetmap.org/copyright). "
        "Landmark details enriched from Wikidata (CC0). Travel guidance from Wikivoyage (CC BY-SA 4.0)."
    ).italic = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
